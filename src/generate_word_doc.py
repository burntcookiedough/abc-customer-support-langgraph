from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTIFACTS = ROOT / "artifacts"
DOCS = ROOT / "docs"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def build_doc() -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI-Powered Customer Support Automation System")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = subtitle.add_run("ABC Technologies | LangGraph Project Documentation")
    sub.font.size = Pt(12)
    sub.font.color.rgb = RGBColor(89, 89, 89)

    add_heading(document, "Project Overview")
    document.add_paragraph(
        "ABC Technologies receives customer support requests for product information, technical issues, "
        "billing queries, account management, and refunds. This project implements a LangGraph workflow "
        "that classifies each query, routes it to the correct department, retrieves company knowledge, "
        "uses SQLite conversation memory, handles high-risk approval, and generates a supervised final response."
    )

    add_heading(document, "Business Requirements Covered")
    add_bullets(
        document,
        [
            "Accept customer queries.",
            "Identify the customer issue type.",
            "Route queries to Sales, Technical Support, Billing, Account, or Memory Recall.",
            "Retrieve relevant context from company policy, pricing, technical manual, and FAQ documents.",
            "Store and recall previous customer interactions using SQLite.",
            "Escalate refund, cancellation, account closure, compensation, and management escalation requests for approval.",
            "Generate final customer responses through a supervisor validation node.",
        ],
    )

    add_heading(document, "LangGraph Workflow")
    document.add_paragraph(
        "The workflow starts with SQLite memory loading, then classifies intent. Department requests use RAG retrieval "
        "before reaching a specialized support agent. High-risk requests pass through human approval. All responses are "
        "validated by the supervisor and stored in SQLite memory."
    )
    diagram = ARTIFACTS / "workflow_diagram.png"
    if diagram.exists():
        document.add_picture(str(diagram), width=Inches(6.8))
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(document, "State Structure")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "State Field", True)
    set_cell_text(table.rows[0].cells[1], "Purpose", True)
    fields = [
        ("customer_id", "Identifies the customer for memory storage and recall."),
        ("query", "Stores the current customer request."),
        ("intent", "Stores the classified route: Sales, Technical Support, Billing, Account, or Memory Recall."),
        ("retrieved_context", "Stores RAG snippets from company documents."),
        ("conversation_history", "Stores recent SQLite memory records."),
        ("approval_required / approval_status", "Tracks human-in-the-loop review for high-risk requests."),
        ("draft_response / final_response", "Stores agent and supervisor-generated responses."),
        ("trace", "Records execution steps for demonstration and debugging."),
    ]
    for field, purpose in fields:
        row = table.add_row().cells
        set_cell_text(row[0], field)
        set_cell_text(row[1], purpose)

    add_heading(document, "Task Coverage")
    task_table = document.add_table(rows=1, cols=3)
    task_table.style = "Table Grid"
    for idx, label in enumerate(["Task", "Requirement", "Implementation"]):
        set_cell_text(task_table.rows[0].cells[idx], label, True)
    tasks = [
        ("1", "Design LangGraph workflow", "Implemented in support_system.py and shown in workflow_diagram.png."),
        ("2", "Create State structure", "SupportState manages customer, query, context, approval, and response fields."),
        ("3", "Intent Classification node", "classify_intent categorizes requests by department and memory recall."),
        ("4", "Conditional routing", "LangGraph conditional edges route requests to the correct node."),
        ("5", "Specialized agents", "Sales, Technical, Billing, and Account support nodes are implemented."),
        ("6", "RAG pipeline", "LocalRAG retrieves context from four company documents in data/."),
        ("7", "SQLite memory", "SQLiteMemory stores and recalls conversations in artifacts/memory.db."),
        ("8", "Human-in-the-loop approval", "High-risk patterns route to the human_approval node."),
        ("9", "Supervisor agent", "supervisor validates and finalizes customer responses."),
        ("10", "Project demonstration", "run_demo.py executes the five required sample queries."),
    ]
    for task, requirement, implementation in tasks:
        row = task_table.add_row().cells
        set_cell_text(row[0], task)
        set_cell_text(row[1], requirement)
        set_cell_text(row[2], implementation)

    add_heading(document, "Knowledge Base Documents")
    add_bullets(
        document,
        [
            "Company Policy Document: refunds, cancellations, account closure, compensation, and escalation rules.",
            "Pricing Guide: Starter, Professional, and Enterprise plan information.",
            "Technical Manual: login, file upload, installation, and configuration guidance.",
            "FAQ Document: password reset, invoices, profile updates, and support hours.",
        ],
    )

    add_heading(document, "Human-in-the-Loop Approval")
    document.add_paragraph(
        "The system marks refund requests, subscription cancellation, account closure, compensation, and management escalation "
        "as high-risk. These requests are not finalized directly by the support agent; they route to the human approval node "
        "before supervisor validation."
    )

    add_heading(document, "SQLite Memory")
    document.add_paragraph(
        "The SQLite database contains a conversations table with customer ID, query, intent, response, and timestamp. "
        "The memory recall query retrieves the customer's most recent previous issue."
    )

    add_heading(document, "Demonstration Results")
    demo_table = document.add_table(rows=1, cols=4)
    demo_table.style = "Table Grid"
    for idx, label in enumerate(["Query", "Expected Path", "Actual Route", "Result"]):
        set_cell_text(demo_table.rows[0].cells[idx], label, True)
    demo_rows = [
        ("What are the pricing plans available for your software?", "Sales", "Sales", "Returned plan information from Pricing Guide."),
        ("I forgot my account password.", "Account", "Account", "Returned password reset guidance from FAQ."),
        ("My application crashes whenever I upload a file.", "Technical Support", "Technical Support", "Returned file upload troubleshooting guidance."),
        ("I need a refund for my annual subscription.", "Billing with approval", "Billing", "Routed through human approval before final response."),
        ("What was my previous support issue?", "Memory recall", "Memory Recall", "Recalled the previous Billing refund issue from SQLite."),
    ]
    for query, expected, actual, result in demo_rows:
        row = demo_table.add_row().cells
        set_cell_text(row[0], query)
        set_cell_text(row[1], expected)
        set_cell_text(row[2], actual)
        set_cell_text(row[3], result)

    add_heading(document, "Submission Files")
    add_bullets(
        document,
        [
            "Source code: src/support_system.py, src/run_demo.py, src/generate_artifacts.py, src/generate_word_doc.py.",
            "README.md: setup steps and run instructions.",
            "Workflow diagram: artifacts/workflow_diagram.png.",
            "Screenshots PDF: artifacts/screenshots.pdf.",
            "SQLite memory: artifacts/memory.db and artifacts/memory_schema.sql.",
            "Demo output: artifacts/demo_output.txt.",
        ],
    )

    output = DOCS / "ABC_Technologies_Customer_Support_Automation_Report.docx"
    document.save(output)
    return output


if __name__ == "__main__":
    print(build_doc())
